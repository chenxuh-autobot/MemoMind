from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path(__file__).resolve().parents[1] / "docs" / "真机闭环验收表.docx"


BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
TEXT = RGBColor(0x11, 0x11, 0x11)
MUTED = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "E8EEF5"
BORDER = "BFCAD8"


@dataclass
class ChecklistRow:
    item_id: str
    check_point: str
    steps: str
    expected: str
    result: str = "□通过  □失败  □阻塞"


SECTIONS: list[tuple[str, str, list[ChecklistRow]]] = [
    (
        "一、验收准备",
        "先确认设备、安装包、模型和基础权限都准备到位，避免把环境问题误判成产品问题。",
        [
            ChecklistRow("P1", "设备版本", "打开系统设置，确认 Android 版本。", "Android 13 及以上。"),
            ChecklistRow("P2", "应用安装", "通过 Android Studio 或 adb 安装 debug 包。", "App 能正常安装并启动。"),
            ChecklistRow("P3", "模型准备", "打开 Home 页，检查本地模型与会话状态。", "模型目录已就绪，会话打开结果显示 success=true。"),
            ChecklistRow("P4", "权限准备", "首次进入 Capture，按提示授权录音与文件读取。", "权限申请链路正常，无异常弹窗。"),
        ],
    ),
    (
        "二、主链路闭环",
        "验证最基本的“文字输入 -> 本地纪要生成 -> 结果查看 -> 历史留存”是否稳定。",
        [
            ChecklistRow("M1", "文字纪要生成", "在 Capture 输入标题和补充文字，点击生成。", "成功跳转到 Result，并看到结构化纪要。"),
            ChecklistRow("M2", "结果完整性", "检查 Result 中的总结、背景、行动项、标签等字段。", "主要字段可展示，无明显空白或错乱。"),
            ChecklistRow("M3", "历史回看", "进入 History 查看最近任务。", "新任务出现在历史列表中。"),
            ChecklistRow("M4", "退出重进", "关闭 App 后重新打开。", "历史任务和最近纪要仍能读取。"),
        ],
    ),
    (
        "三、图片链路闭环",
        "验证图片输入是否能经过 OCR 和图片语义摘要，再进入本地纪要生成。",
        [
            ChecklistRow("I1", "图片选择", "在 Capture 选择一张含中文文字的图片。", "图片文件引用显示正常。"),
            ChecklistRow("I2", "图片 OCR", "点击“从图片识别 OCR”。", "OCR 文本自动回填。"),
            ChecklistRow("I3", "图片要点", "点击“提取图片要点”。", "图片内容补充自动回填语义摘要。"),
            ChecklistRow("I4", "图片纪要生成", "结合图片结果提交纪要任务。", "Result 能反映图片来源内容。"),
        ],
    ),
    (
        "四、语音链路闭环",
        "验证实时转写、自动录音留档和结果页回放是否可用。",
        [
            ChecklistRow("A1", "麦克风转写", "点击“开始麦克风转写”，说一段 10 到 20 秒中文。", "录音转写文本持续或最终回填。"),
            ChecklistRow("A2", "停止转写", "点击“停止麦克风转写”。", "状态恢复正常，并生成录音素材引用。"),
            ChecklistRow("A3", "语音纪要生成", "提交当前任务。", "结构化纪要成功生成，任务写入本地。"),
            ChecklistRow("A4", "录音回放", "在 Result 页点击播放录音。", "可正常播放和停止播放。"),
        ],
    ),
    (
        "五、音频文件级重跑转写",
        "验证已保存录音可以再次做文件级转写，并将结果重新带回 Capture 页面。",
        [
            ChecklistRow("R1", "触发重跑转写", "在 Result 的 AUDIO 素材卡片点击“重跑转写到 Capture”。", "按钮可点击，并进入处理状态。"),
            ChecklistRow("R2", "回填结果", "等待处理完成。", "自动跳回 Capture，并把新转写结果回填到录音转写文本。"),
            ChecklistRow("R3", "素材保留", "检查当前草稿中的录音素材引用。", "原录音素材仍保留，可继续提交。"),
            ChecklistRow("R4", "再次提交", "基于重跑结果再次生成纪要。", "可形成第二轮纪要结果。"),
        ],
    ),
    (
        "六、异常与稳定性观察",
        "记录真机体验层面的实际表现，便于后续优化答辩内容和工程路线。",
        [
            ChecklistRow("S1", "响应耗时", "记录 OCR、纪要生成、重跑转写的耗时感受。", "耗时可接受，或能明确记录瓶颈。"),
            ChecklistRow("S2", "发热与卡顿", "连续执行多条任务。", "无明显卡死；若有发热或掉帧需记录。"),
            ChecklistRow("S3", "错误提示", "故意在空输入、权限拒绝等场景操作。", "有可理解的状态提示，不是静默失败。"),
            ChecklistRow("S4", "可演示性", "按比赛答辩节奏完整演示一遍。", "可在有限时间内顺利讲清链路。"),
        ],
    ),
]


def set_document_geometry(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def set_run_font(run, *, size: int, color: RGBColor = TEXT, bold: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def apply_paragraph_format(paragraph, *, before: int, after: int, line: float, alignment=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def ensure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    h1 = document.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h1.font.size = Pt(16)
    h1.font.color.rgb = BLUE
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing = 1.25

    h2 = document.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h2.font.size = Pt(13)
    h2.font.color.rgb = BLUE
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.line_spacing = 1.25


def set_cell_margins(cell, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        element = tc_mar.find(qn(f"w:{key}"))
        if element is None:
            element = OxmlElement(f"w:{key}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER.rgb.hex if hasattr(BORDER, "hex") else "BFCAD8")


def set_table_layout_fixed(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def add_title_block(document: Document) -> None:
    title = document.add_paragraph()
    apply_paragraph_format(title, before=0, after=4, line=1.0)
    title_run = title.add_run("Creative AI Android 真机闭环验收表")
    set_run_font(title_run, size=22, color=TEXT, bold=True)

    subtitle = document.add_paragraph()
    apply_paragraph_format(subtitle, before=0, after=12, line=1.15)
    subtitle_run = subtitle.add_run("适用范围：本地 Qwen + MNN 纪要链路、图片 OCR/语义摘要、麦克风转写、音频文件级重跑转写。")
    set_run_font(subtitle_run, size=11, color=MUTED, bold=False)

    meta = document.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.autofit = False
    set_table_layout_fixed(meta)
    widths = (Inches(1.875), Inches(4.625))
    meta_rows = [
        ("验收目标", "确认项目已经达到可演示闭环，并记录真实设备上的稳定性与异常表现。"),
        ("建议设备", "Android 13 及以上真机，优先使用支持本地语音识别服务的设备。"),
        ("建议执行顺序", "准备 -> 主链路 -> 图片 -> 语音 -> 重跑转写 -> 持久化 -> 异常观察"),
        ("验收结论", "□ 通过    □ 有条件通过    □ 未通过"),
    ]
    for row, (label, value) in zip(meta.rows, meta_rows):
        row.cells[0].width = widths[0]
        row.cells[1].width = widths[1]
        for idx, text in enumerate((label, value)):
            cell = row.cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            apply_paragraph_format(paragraph, before=0, after=3, line=1.15)
            run = paragraph.add_run(text)
            set_run_font(run, size=11, color=TEXT, bold=idx == 0)
    shade_cell(meta.rows[0].cells[0], HEADER_FILL)
    shade_cell(meta.rows[1].cells[0], HEADER_FILL)
    shade_cell(meta.rows[2].cells[0], HEADER_FILL)
    shade_cell(meta.rows[3].cells[0], HEADER_FILL)
    set_table_borders(meta)


def add_section_table(document: Document, title: str, intro: str, rows: list[ChecklistRow]) -> None:
    heading = document.add_paragraph(style="Heading 1")
    heading.add_run(title)

    intro_paragraph = document.add_paragraph()
    apply_paragraph_format(intro_paragraph, before=0, after=8, line=1.25)
    intro_run = intro_paragraph.add_run(intro)
    set_run_font(intro_run, size=11, color=MUTED)

    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_layout_fixed(table)
    header = table.rows[0].cells
    headers = ["编号", "验收点", "操作步骤", "通过标准", "结果"]
    widths = [Inches(0.55), Inches(1.35), Inches(2.15), Inches(1.45), Inches(1.0)]
    for cell, text, width in zip(header, headers, widths):
        cell.width = width
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        apply_paragraph_format(p, before=0, after=2, line=1.1, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run(text)
        set_run_font(run, size=10, color=DARK_BLUE, bold=True)

    for row_data in rows:
        row = table.add_row().cells
        values = [row_data.item_id, row_data.check_point, row_data.steps, row_data.expected, row_data.result]
        for idx, (cell, text, width) in enumerate(zip(row, values, widths)):
            cell.width = width
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 4) else WD_ALIGN_PARAGRAPH.LEFT
            apply_paragraph_format(p, before=0, after=3, line=1.15, alignment=align)
            run = p.add_run(text)
            set_run_font(run, size=10 if idx != 4 else 9, color=TEXT, bold=idx == 0)

    set_table_borders(table)


def add_issue_log(document: Document) -> None:
    heading = document.add_paragraph(style="Heading 1")
    heading.add_run("七、问题记录与复测")

    intro = document.add_paragraph()
    apply_paragraph_format(intro, before=0, after=8, line=1.25)
    run = intro.add_run("如果现场出现阻塞、错误提示、结果异常或性能问题，请按下面表格记录，方便后续复测和答辩说明。")
    set_run_font(run, size=11, color=MUTED)

    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_layout_fixed(table)
    widths = [Inches(0.8), Inches(1.5), Inches(1.4), Inches(1.6), Inches(1.2)]
    headers = ["序号", "问题场景", "现象描述", "临时结论", "是否复测"]
    for cell, text, width in zip(table.rows[0].cells, headers, widths):
        cell.width = width
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        apply_paragraph_format(p, before=0, after=2, line=1.1, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run(text)
        set_run_font(run, size=10, color=DARK_BLUE, bold=True)

    for index in range(1, 6):
        row = table.add_row().cells
        values = [str(index), "", "", "", "□是  □否"]
        for idx, (cell, value, width) in enumerate(zip(row, values, widths)):
            cell.width = width
            set_cell_margins(cell, top=120, bottom=120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            apply_paragraph_format(p, before=0, after=2, line=1.15, alignment=WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 4) else WD_ALIGN_PARAGRAPH.LEFT)
            run = p.add_run(value)
            set_run_font(run, size=10, color=TEXT)

    set_table_borders(table)


def add_footer(document: Document) -> None:
    section = document.sections[0]
    footer = section.footer.paragraphs[0]
    apply_paragraph_format(footer, before=0, after=0, line=1.0, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    footer_run = footer.add_run("Creative AI Android | 真机闭环验收")
    set_run_font(footer_run, size=9, color=MUTED)


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    set_document_geometry(document)
    ensure_styles(document)
    add_title_block(document)

    for title, intro, rows in SECTIONS:
        add_section_table(document, title, intro, rows)

    add_issue_log(document)
    add_footer(document)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
